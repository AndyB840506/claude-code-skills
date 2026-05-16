from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_colored_heading(doc, text, level=1, color=(255, 140, 0)):
    """Add heading with custom color"""
    heading = doc.add_heading(text, level=level)
    heading_format = heading.paragraph_format
    heading_format.space_before = Pt(12)
    heading_format.space_after = Pt(12)

    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)

    return heading

def shade_cell(cell, fill_color):
    """Shade table cell with color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# HEADER
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("MR. PUTRID'S DEN")
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(26, 26, 26)

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline_run = tagline.add_run("Donde los riffs encuentran el whisky")
tagline_run.font.size = Pt(16)
tagline_run.font.italic = True
tagline_run.font.color.rgb = RGBColor(255, 215, 0)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("Propuesta de Podcast — Conversaciones, Bandas, Eventos Underground")
subtitle_run.font.size = Pt(11)
subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()  # Spacing

# PERFIL DEL PODCAST
add_colored_heading(doc, "Perfil del Podcast", level=1)

# Cards table
table = doc.add_table(rows=2, cols=2)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
headers = ["📻 Formato", "⏱️ Duración & Cadencia"]
for i, header_text in enumerate(headers):
    cell = header_cells[i]
    shade_cell(cell, "FF8C00")
    p = cell.paragraphs[0]
    p_run = p.add_run(header_text)
    p_run.font.bold = True
    p_run.font.color.rgb = RGBColor(255, 255, 255)

# Content row
content_cells = table.rows[1].cells
content_texts = [
    "Co-host\nAndrés & Juan — conversación natural sobre música, bandas y eventos underground",
    "60 minutos | Semanal\nUn episodio nuevo cada semana"
]
for i, content_text in enumerate(content_texts):
    cell = content_cells[i]
    p = cell.paragraphs[0]
    p.text = content_text
    p.paragraph_format.line_spacing = 1.15

# Add more info cards
p = doc.add_paragraph()
p_run = p.add_run("🎯 Audiencia: ")
p_run.font.bold = True
p.add_run("Bogotanos apasionados por la música. Rock, metal, jazz, géneros variados (20-45+ años)")

p = doc.add_paragraph()
p_run = p.add_run("📱 Plataformas: ")
p_run.font.bold = True
p.add_run("Spotify + RSS (distribución a otras plataformas)")

# Description
add_colored_heading(doc, "Descripción", level=3)
description = doc.add_paragraph(
    "Mr. Putrid's Den es un podcast conversacional co-hosted por Andrés y Juan, un espacio dedicado a explorar rock, metal, jazz y múltiples géneros musicales. Con Juan como promotor de eventos underground, el podcast invita artistas emergentes, comparte historias de conciertos únicos, y ofrece análisis profundo de bandas y álbumes icónicos. Un lugar donde la música, la conversación natural y la pasión por los riffs se encuentran."
)

# IDENTIDAD VISUAL
doc.add_paragraph()
add_colored_heading(doc, "Identidad Visual", level=1)

add_colored_heading(doc, "Paleta de Colores", level=3)
color_table = doc.add_table(rows=1, cols=3)
color_cells = color_table.rows[0].cells

colors_data = [
    ("#1A1A1A", "Negro Principal"),
    ("#FF8C00", "Naranja Cálido"),
    ("#FFD700", "Dorado")
]

for i, (hex_color, name) in enumerate(colors_data):
    cell = color_cells[i]
    shade_cell(cell, hex_color.replace("#", ""))
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_run = p.add_run(name)
    p_run.font.color.rgb = RGBColor(255, 255, 255) if i != 2 else RGBColor(26, 26, 26)
    p_run.font.bold = True

add_colored_heading(doc, "Estilo Visual", level=3)
p = doc.add_paragraph()
p_run = p.add_run("Vintage collage + Pop art")
p_run.font.bold = True
doc.add_paragraph("Referencias a bandas icónicas, estética underground, conciertos, vinilos. Acentos de neon verde y magenta. Tono: nostálgico pero energético, profesional pero accesible.")

add_colored_heading(doc, "Audio Branding", level=3)
p = doc.add_paragraph()
p_run = p.add_run("🎸 INTRO (30 seg)")
p_run.font.bold = True
p.add_run("\nStevie Ray Vaughn meets Black Label Society. Guitarra blues eléctrica + heavy metal con voz femenina gutural.")

p = doc.add_paragraph()
p_run = p.add_run("🎷 OUTRO (30 seg)")
p_run.font.bold = True
p.add_run("\nJazz 1950s estilo Billie Holiday. Cierre elegante con trio de mujeres. 'The den, it's still warm, a shelter from the storm...'")

# GUÍA DE TONO
doc.add_paragraph()
add_colored_heading(doc, "Guía de Tono", level=1)

add_colored_heading(doc, "5 Reglas de Escritura", level=3)

rules = [
    ("Lenguaje bogotano/colombiano neutro", "Usa: 'parcero/a', 'hermano/a', 'bacano', 'qué onda', 'qué vueltas'. Evita: paisajismos, mexicanismos, regionalismos de otras zonas."),
    ("Conversación natural con estructura", "El script es guía, no lectura. Lanza premisas, datos, anécdotas, y deja que fluya entre Andrés y Juan."),
    ("Respeta a los artistas y eventos", "Cuando hables de bandas, artistas emergentes o eventos underground, hazlo con admiración y conocimiento."),
    ("Datos + opinión + anécdota", "Cada episodio debe tener: un dato sorprendente, análisis profundo, y una historia personal de concierto o música."),
    ("Whisky, riffs y conversación", "El podcast es sobre tres cosas: la música, la amistad, y el disfrute. Mantén ese balance.")
]

for i, (title_rule, description_rule) in enumerate(rules, 1):
    p = doc.add_paragraph(f"{i}. ", style='List Number')
    p_run = p.add_run(title_rule)
    p_run.font.bold = True
    p.add_run(f"\n{description_rule}")

# PILARES DE CONTENIDO
doc.add_paragraph()
add_colored_heading(doc, "Pilares de Contenido", level=1)
doc.add_paragraph("Estos 8 pilares rotan para mantener variedad y estructura en los episodios:")

pillars = [
    ("🎵 Análisis Profundo", "Desglosar una banda, álbum o canción icónica."),
    ("🏆 Top 5 / Rankings", "Ranking de riffs, canciones, albums."),
    ("📚 Historia de Género", "Evolución de rock, metal, jazz, fusion."),
    ("🎤 Debates Clásicos", "¿Beatles vs Led Zeppelin? Discusión apasionada."),
    ("🎬 Historias de Conciertos", "Anécdotas personales y momentos memorables."),
    ("👥 Invitados Artistas", "Entrevistas con artistas emergentes."),
    ("🌍 Spotlight: Eventos", "Events underground, conciertos próximos."),
    ("🎧 Reacciones a Música", "Escuchamos música nueva en vivo.")
]

for title_p, desc_p in pillars:
    p = doc.add_paragraph()
    p_run = p.add_run(title_p)
    p_run.font.bold = True
    p.add_run(f" — {desc_p}")

# ROADMAP
doc.add_paragraph()
add_colored_heading(doc, "Roadmap — 10 Episodios", level=1)

roadmap_table = doc.add_table(rows=11, cols=4)
roadmap_table.style = 'Light Grid Accent 1'

# Header
header_cells = roadmap_table.rows[0].cells
headers = ["EP", "Título", "Pilar", "Tipo"]
for i, header_text in enumerate(headers):
    cell = header_cells[i]
    shade_cell(cell, "1A1A1A")
    p = cell.paragraphs[0]
    p_run = p.add_run(header_text)
    p_run.font.bold = True
    p_run.font.color.rgb = RGBColor(255, 255, 255)

# Data
roadmap_data = [
    ("001", "Bienvenidos a Mr. Putrid's Den", "Presentación", "Co-host"),
    ("002", "Black Sabbath: El génesis del heavy metal", "Análisis profundo", "Co-host"),
    ("003", "Top 5: Los mejores riffs de metal", "Top 5 / Ranking", "Co-host"),
    ("004", "Jazz fusion: De Herbie Hancock a hoy", "Historia de género", "Co-host"),
    ("005", "¿The Beatles o Led Zeppelin?", "Debate clásico", "Co-host"),
    ("006", "Stories: Conciertos únicos que nos marcaron", "Historias", "Co-host"),
    ("007", "Entrevista: Artista emergente", "Invitado artista", "Co-host + Guest"),
    ("008", "Spotlight: Eventos underground", "Info conciertos", "Co-host"),
    ("009", "Rock progresivo: Yes & Genesis", "Historia de género", "Co-host"),
    ("010", "New music reaction", "Reacciones", "Co-host"),
]

for row_idx, (ep, titulo, pilar, tipo) in enumerate(roadmap_data, 1):
    cells = roadmap_table.rows[row_idx].cells
    cells[0].text = ep
    cells[1].text = titulo
    cells[2].text = pilar
    cells[3].text = tipo

# FOOTER
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run("Mr. Putrid's Den — Propuesta de Podcast")
footer_run.font.bold = True
footer_run.font.size = Pt(10)

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer2_run = footer2.add_run("Generado: 2026-05-15 | Donde los riffs encuentran el whisky 🥃🎸")
footer2_run.font.size = Pt(9)
footer2_run.font.color.rgb = RGBColor(150, 150, 150)

# Save
output_path = r"e:\Claude Project\Claude Projects\kit-skill-creator\other\kit-podcast-creator\documents\Mr-Putrids-Den-Propuesta.docx"
doc.save(output_path)
print("Documento Word creado correctamente")
